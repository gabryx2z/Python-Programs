import datetime
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Definisci il nome del file in cui registrare gli accessi e le uscite
file_name = "accessi_uscite.docx"

# Verifica se il file esiste già. Se sì, apri il documento esistente, altrimenti crea un nuovo documento
if os.path.exists(file_name):
    document = Document(file_name)
else:
    document = Document()

# Chiedi all'utente se l'azione è un'entrata o un'uscita
azione = input("L'azione è un'entrata o un'uscita? ")

# Ottieni la data e l'ora correnti
now = datetime.datetime.now()

# Aggiungi un nuovo paragrafo al documento
p = document.add_paragraph()

# Aggiungi la data e l'ora correnti al paragrafo
p.add_run(now.strftime("%Y-%m-%d %H:%M:%S ")).bold = True

# Aggiungi l'azione al paragrafo
p.add_run(azione)

# Allinea il paragrafo a sinistra
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Salva il documento
document.save(file_name)

# Stampa il contenuto del documento
print("Accessi e uscite registrati:")
for paragraph in document.paragraphs:
    print(paragraph.text)
#Il programma utilizza la libreria datetime per ottenere la data e l'ora correnti, la libreria os
#per verificare se il file esiste già, e la libreria python-docx per creare il file Word e
#aggiungere i paragrafi. Il programma chiede all'utente se l'azione è un'entrata o un'uscita,
#registra la data e l'ora correnti insieme all'azione in un nuovo paragrafo nel file Word, e alla
#fine stampa il contenuto del documento.
    
